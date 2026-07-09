from docx import Document


# =====================================================
# EXTRACT TEXT FROM DOCX
# =====================================================

def extract_docx_text(uploaded_file):
    """
    Extract all readable text from a DOCX file.
    """

    try:

        uploaded_file.seek(0)

        document = Document(uploaded_file)

        text = []

        for paragraph in document.paragraphs:

            if paragraph.text.strip():

                text.append(paragraph.text)

        return "\n".join(text)

    except Exception as e:

        return f"Error reading DOCX: {e}"


# =====================================================
# EXTRACT TABLES
# =====================================================

def extract_docx_tables(uploaded_file):
    """
    Extract all tables from a DOCX document.
    Returns a list of tables.
    """

    try:

        uploaded_file.seek(0)

        document = Document(uploaded_file)

        tables = []

        for table in document.tables:

            rows = []

            for row in table.rows:

                rows.append(

                    [

                        cell.text.strip()

                        for cell in row.cells

                    ]

                )

            tables.append(rows)

        return tables

    except Exception:

        return []


# =====================================================
# DOCUMENT INFORMATION
# =====================================================

def get_docx_information(uploaded_file):
    """
    Returns useful metadata about a DOCX document.
    """

    try:

        uploaded_file.seek(0)

        document = Document(uploaded_file)

        info = {

            "Paragraphs": len(document.paragraphs),

            "Tables": len(document.tables),

            "Sections": len(document.sections)

        }

        return info

    except Exception:

        return None


# =====================================================
# WORD COUNT
# =====================================================

def document_word_count(text):

    if not text:

        return 0

    return len(text.split())


# =====================================================
# CHARACTER COUNT
# =====================================================

def document_character_count(text):

    if not text:

        return 0

    return len(text)


# =====================================================
# PREVIEW
# =====================================================

def preview_text(text, characters=3000):

    if not text:

        return ""

    return text[:characters]